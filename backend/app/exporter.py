"""
Export-Funktionen für Jobs und Exporte.

Dieser Modul stellt die Logik bereit, um aktuelle Artefaktinhalte aus
der Datenbank in unterschiedliche Dateiformate zu exportieren und als
ZIP-Archiv zu bündeln. Es unterstützt Text (``.txt``), Markdown (``.md``),
Word (``.docx``) und PDF (``.pdf``) und sorgt dabei für eine einfache,
lesbare Struktur der erzeugten Dokumente.

Die PDF-Erzeugung nutzt ``reportlab`` mit symmetrischen Seitenrändern
(50 pt links und rechts) und einfachem Zeilenumbruch. Die DOCX-Erzeugung
wandelt Markdown-Überschriften in Word-Überschriften (Heading 1–3),
Zahlenlisten in nummerierte Listen und Aufzählungen in Bullet-Listen.

Alle erzeugten Dateien werden in ein temporäres Verzeichnis geschrieben
und anschließend in ein ZIP-Archiv gepackt. Temporäre Dateien werden
anschließend entfernt.

Block 09 – hier ergänzt und verbessert in Block 10: Layout und
Dokumentenstruktur.
"""

from __future__ import annotations

import os
import re
import zipfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import List, Tuple

from sqlalchemy import select
from sqlalchemy.orm import Session

from .models import Artifact, ArtifactVersion
from .settings import EXPORT_DIR

# Optional Abhängigkeiten für DOCX und PDF
from docx import Document  # type: ignore
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer  # type: ignore
from reportlab.lib.styles import getSampleStyleSheet  # type: ignore
from reportlab.lib.pagesizes import A4  # type: ignore


@dataclass(frozen=True)
class ExportItem:
    """Container für ein zu exportierendes Artefakt.

    ``artifact_id`` ist die UUID des Artefakts. ``filename_base`` dient
    als Basisname der Zieldatei ohne Extension. ``content_md`` enthält
    den Markdown-Inhalt des aktuellen Artefaktstands.
    """

    artifact_id: str
    filename_base: str
    content_md: str


def _safe_filename(name: str) -> str:
    """Säubert einen Dateinamen, um gefährliche Zeichen zu entfernen.

    Es werden nur alphanumerische Zeichen sowie Punkt, Unterstrich und
    Bindestrich zugelassen. Gleichzeitig wird die Länge begrenzt, um
    Dateisystemprobleme zu vermeiden. Falls der Name leer ist, wird
    ``artifact`` zurückgegeben.
    """

    s = name.strip()
    s = re.sub(r"[^a-zA-Z0-9._-]+", "_", s)
    s = s.strip("._-")
    return s[:120] if s else "artifact"


def _load_artifacts_current(db: Session, artifact_ids: List[str]) -> List[ExportItem]:
    """Lädt die aktuellen Versionen der angegebenen Artefakte.

    Gibt eine Liste von ``ExportItem``-Instanzen zurück. Für nicht
    existierende Artefakte wird ein Platzhalter mit Hinweistext
    erstellt, damit der Export konsistent bleibt.
    """

    items: List[ExportItem] = []
    if not artifact_ids:
        return items

    arts = db.execute(select(Artifact).where(Artifact.id.in_(artifact_ids))).scalars().all()
    art_by_id = {a.id: a for a in arts}

    for aid in artifact_ids:
        art = art_by_id.get(aid)
        if not art:
            items.append(
                ExportItem(
                    artifact_id=aid,
                    filename_base=_safe_filename(aid),
                    content_md=f"# Fehlender Datensatz\n\nArtefakt {aid} wurde nicht gefunden.\n",
                )
            )
            continue

        v = db.execute(
            select(ArtifactVersion)
            .where(ArtifactVersion.artifact_id == art.id)
            .where(ArtifactVersion.version == art.current_version)
        ).scalars().first()
        content = v.content_md if v else ""
        base = _safe_filename(
            f"{art.type}_{getattr(art, 'title', '')}" if getattr(art, "title", None) else art.type
        )
        items.append(ExportItem(artifact_id=aid, filename_base=base, content_md=content or ""))

    return items


def _write_txt(path: Path, content: str) -> None:
    path.write_text(content, encoding="utf-8")


def _write_md(path: Path, content: str) -> None:
    path.write_text(content, encoding="utf-8")


def _write_docx(path: Path, content_md: str) -> None:
    """Schreibt den Markdown-Inhalt in eine DOCX-Datei mit einfacher Struktur.

    Es werden Überschriften der Stufen 1–3 erkannt und in die
    entsprechenden Word-Heading-Levels umgesetzt. Numerierte Zeilen
    werden als nummerierte Liste ("List Number") formatiert, und
    Aufzählungen mit ``-`` oder ``*`` werden als Bullets ("List Bullet")
    umgesetzt. Alle anderen Zeilen werden als normale Absätze
    geschrieben.
    """

    doc = Document()
    lines = content_md.splitlines()
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif re.match(r"^\s*\d+\.\s+", line):
            text = re.sub(r"^\s*\d+\.\s+", "", line).strip()
            doc.add_paragraph(text, style="List Number")
        elif re.match(r"^\s*[-*]\s+", line):
            text = re.sub(r"^\s*[-*]\s+", "", line).strip()
            doc.add_paragraph(text, style="List Bullet")
        else:
            doc.add_paragraph(line)
    doc.save(str(path))


def _write_pdf(path: Path, content_md: str) -> None:
    """Schreibt den Markdown-Inhalt in eine PDF-Datei.

    Dabei werden Überschriften erkannt und mit vordefinierten
    Layout-Stilen versehen. Listenpunkte werden durch ein "•" ersetzt.
    Es werden symmetrische Seitenränder verwendet und automatische
    Zeilenumbrüche über das ReportLab-Framework erzeugt, um das
    Überlaufen des Textes zu verhindern.
    """

    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=50,
        rightMargin=50,
        topMargin=50,
        bottomMargin=50,
    )
    styles = getSampleStyleSheet()
    flow: List = []
    lines = content_md.splitlines()
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            # Leerzeile → Abstand
            flow.append(Spacer(1, 8))
            continue
        if line.startswith("### "):
            text = line[4:].strip()
            flow.append(Paragraph(text, styles.get("Heading3", styles["Heading2"])) )
        elif line.startswith("## "):
            text = line[3:].strip()
            flow.append(Paragraph(text, styles.get("Heading2", styles["Heading2"])) )
        elif line.startswith("# "):
            text = line[2:].strip()
            flow.append(Paragraph(text, styles.get("Heading1", styles["Heading1"])) )
        elif re.match(r"^\s*\d+\.\s+", line):
            # Numerierte Liste → Bullet für PDF
            text = re.sub(r"^\s*\d+\.\s+", "", line).strip()
            flow.append(Paragraph(f"• {text}", styles["Normal"]))
        elif re.match(r"^\s*[-*]\s+", line):
            text = re.sub(r"^\s*[-*]\s+", "", line).strip()
            flow.append(Paragraph(f"• {text}", styles["Normal"]))
        else:
            flow.append(Paragraph(line, styles["Normal"]))
        # Kleiner Abstand zwischen den Abschnitten
        flow.append(Spacer(1, 4))
    doc.build(flow)


def export_artifacts_to_zip(
    db: Session,
    artifact_ids: List[str],
    export_format: str,
    job_id: str,
) -> Tuple[str, str]:
    """Exportiert Artefakte in das gewünschte Format und packt sie in ein ZIP.

    ``export_format`` kann ``txt``, ``md``, ``docx`` oder ``pdf`` sein.
    Das ZIP-Archiv wird im ``EXPORT_DIR`` gespeichert und enthält für
    jedes Artefakt eine Datei. Es wird ein Tupel aus Dateiname des
    ZIPs und dem absoluten Pfad zurückgegeben.
    """

    fmt = (export_format or "md").lower().strip()
    if fmt not in {"txt", "md", "docx", "pdf"}:
        fmt = "md"

    base_dir = Path(EXPORT_DIR)
    base_dir.mkdir(parents=True, exist_ok=True)

    tmp_dir = base_dir / f"tmp_{job_id}"
    tmp_dir.mkdir(parents=True, exist_ok=True)

    items = _load_artifacts_current(db, artifact_ids)

    if not items:
        note = tmp_dir / "README.txt"
        _write_txt(note, "Kein Export-Inhalt: artifact_ids war leer.\n")

    for it in items:
        out_path = tmp_dir / f"{it.filename_base}.{fmt}"
        if fmt == "txt":
            _write_txt(out_path, it.content_md)
        elif fmt == "md":
            _write_md(out_path, it.content_md)
        elif fmt == "docx":
            _write_docx(out_path, it.content_md)
        elif fmt == "pdf":
            _write_pdf(out_path, it.content_md)

    zip_filename = f"{job_id}.zip"
    zip_path = base_dir / zip_filename
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for p in tmp_dir.iterdir():
            if p.is_file():
                zf.write(p, arcname=p.name)

    # Aufräumen des temporären Verzeichnisses
    for p in tmp_dir.iterdir():
        try:
            p.unlink()
        except Exception:
            pass
    try:
        tmp_dir.rmdir()
    except Exception:
        pass

    return zip_filename, str(zip_path)